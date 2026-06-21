import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_UNDERLINE

# Set paths
HERE = Path(__file__).parent
OUTPUTS_DIR = HERE / "outputs"
DOCS_DIR = HERE / "word_documents"

# Ensure docs directory exists
DOCS_DIR.mkdir(exist_ok=True)

def main():
    print("Scanning outputs folder for JSON files...")
    json_files = list(OUTPUTS_DIR.glob("*.json"))
    
    if not json_files:
        print("No JSON files found in outputs/ directory.")
        return
        
    # Initialize a single Word Document
    doc = Document()
    
    # Keep track of which player types we've already added
    seen_player_types = set()
    examples_added = 0
    
    for json_file in json_files:
        # Load JSON
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
            
        player_type = data.get("player_type", "Unknown Player")
        
        # Only add one example per player type
        if player_type in seen_player_types:
            continue
            
        # Mark this player type as seen
        seen_player_types.add(player_type)
        examples_added += 1
        
        scenario_data = data.get("scenario", {})
        
        # ---------------------------------------------------------
        # Format the entry in the Word Doc
        # ---------------------------------------------------------
        
        # Title (Player Type)
        title_run = doc.add_paragraph().add_run(player_type)
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        # Culture
        culture_p = doc.add_paragraph()
        culture_p.add_run("Culture: ").bold = False
        culture_p.add_run(scenario_data.get("culture", ""))
        
        # Demographics
        doc.add_paragraph("Demographics:")
        doc.add_paragraph(scenario_data.get("demographics", ""))
        
        # Cultural Context
        doc.add_paragraph("Cultural Context:")
        doc.add_paragraph(scenario_data.get("cultural_context", ""))
        
        # Geographic Setting
        doc.add_paragraph("Geographic Setting:")
        doc.add_paragraph(scenario_data.get("geographic_setting", ""))
        
        # Scenario Header
        scenario_header_p = doc.add_paragraph()
        scenario_header_run = scenario_header_p.add_run("SCENARIO")
        scenario_header_run.bold = True
        scenario_header_run.underline = WD_UNDERLINE.SINGLE
        
        # Scenario Text
        scenario_text = scenario_data.get("scenario_text", "")
        paragraphs = scenario_text.split('\n\n')
        
        for p in paragraphs:
            if p.strip():
                doc.add_paragraph(p.strip())
                
        # Add a page break between different player types (except the last one)
        doc.add_page_break()

    # Save the merged document
    output_filename = "Merged_Player_Type_Examples.docx"
    output_path = DOCS_DIR / output_filename
    doc.save(output_path)
    
    print(f"\nSuccessfully generated '{output_filename}'!")
    print(f"It contains {examples_added} unique player type examples.")

if __name__ == "__main__":
    main()
